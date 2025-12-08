"""
Word文档解析模块：提取Word文档中的文本、表格和图片信息
"""
import json
from pathlib import Path
from typing import Dict, List, Any
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
import logging

logger = logging.getLogger(__name__)


class DocxParser:
    """Word文档解析器"""
    
    def __init__(self):
        """初始化解析器"""
        pass
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        解析Word文档
        
        Args:
            file_path: Word文档路径
            
        Returns:
            包含文档结构和内容的字典
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        if file_path.suffix.lower() not in ['.docx', '.doc']:
            raise ValueError(f"不支持的文件格式: {file_path.suffix}")
        
        logger.info(f"解析Word文档: {file_path.name}")
        
        result = {
            "file_name": file_path.name,
            "file_path": str(file_path),
            "content": "",
            "structure": [],
            "paragraphs": [],
            "tables": [],
            "images": [],
            "metadata": {}
        }
        
        try:
            doc = Document(file_path)
            
            # 提取文档属性
            result["metadata"] = self._extract_metadata(doc)
            
            # 解析文档内容
            content_parts = []
            
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # 段落
                    paragraph = Paragraph(element, doc)
                    para_data = self._parse_paragraph(paragraph)
                    
                    if para_data["text"].strip():
                        result["paragraphs"].append(para_data)
                        result["structure"].append({
                            "type": "paragraph",
                            "index": len(result["paragraphs"]) - 1
                        })
                        content_parts.append(para_data["text"])
                
                elif isinstance(element, CT_Tbl):
                    # 表格
                    table = Table(element, doc)
                    table_data = self._parse_table(table)
                    
                    result["tables"].append(table_data)
                    result["structure"].append({
                        "type": "table",
                        "index": len(result["tables"]) - 1
                    })
                    content_parts.append(table_data["text_content"])
            
            # 提取图片信息
            result["images"] = self._extract_images_info(doc)
            
            # 合并所有文本内容
            result["content"] = "\n\n".join(content_parts)
            
            logger.info(f"文档解析完成: {len(result['paragraphs'])} 段落, "
                       f"{len(result['tables'])} 表格, {len(result['images'])} 图片")
        
        except Exception as e:
            logger.error(f"解析文档失败: {str(e)}")
            result["error"] = str(e)
        
        return result
    
    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """提取文档元数据"""
        core_props = doc.core_properties
        
        metadata = {
            "author": core_props.author,
            "title": core_props.title,
            "subject": core_props.subject,
            "created": str(core_props.created) if core_props.created else None,
            "modified": str(core_props.modified) if core_props.modified else None,
            "last_modified_by": core_props.last_modified_by,
        }
        
        return metadata
    
    def _parse_paragraph(self, paragraph: Paragraph) -> Dict[str, Any]:
        """解析段落"""
        para_data = {
            "text": paragraph.text,
            "style": paragraph.style.name if paragraph.style else None,
            "alignment": str(paragraph.alignment) if paragraph.alignment else None,
            "runs": []
        }
        
        # 解析段落中的runs（保留格式信息）
        for run in paragraph.runs:
            run_data = {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_size": run.font.size.pt if run.font.size else None,
                "font_name": run.font.name
            }
            para_data["runs"].append(run_data)
        
        return para_data
    
    def _parse_table(self, table: Table) -> Dict[str, Any]:
        """解析表格"""
        table_data = {
            "rows": len(table.rows),
            "columns": len(table.columns),
            "data": [],
            "text_content": ""
        }
        
        text_parts = []
        
        for i, row in enumerate(table.rows):
            row_data = []
            row_text = []
            
            for j, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                row_data.append({
                    "row": i,
                    "column": j,
                    "text": cell_text
                })
                row_text.append(cell_text)
            
            table_data["data"].append(row_data)
            text_parts.append(" | ".join(row_text))
        
        table_data["text_content"] = "\n".join(text_parts)
        
        return table_data
    
    def _extract_images_info(self, doc: Document) -> List[Dict[str, Any]]:
        """提取文档中的图片信息"""
        images = []
        
        try:
            # 获取文档中的所有图片关系
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_info = {
                        "id": rel.rId,
                        "type": rel.target_ref.split('/')[-1],
                        "content_type": rel.target_part.content_type if hasattr(rel, 'target_part') else None
                    }
                    images.append(image_info)
        
        except Exception as e:
            logger.warning(f"提取图片信息失败: {str(e)}")
        
        return images
    
    def extract_text_only(self, file_path: str) -> str:
        """
        仅提取文档的纯文本内容
        
        Args:
            file_path: Word文档路径
            
        Returns:
            纯文本内容
        """
        doc = Document(file_path)
        
        text_parts = []
        
        # 提取所有段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # 提取所有表格
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    
    def save_result(self, result: Dict, output_path: str):
        """
        保存解析结果到JSON文件
        
        Args:
            result: 解析结果字典
            output_path: 输出文件路径
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        
        logger.info(f"解析结果已保存到: {output_path}")
    
    def compare_structure(self, doc_result: Dict, ocr_results: List[Dict]) -> Dict[str, Any]:
        """
        比较Word文档和OCR结果的结构
        
        Args:
            doc_result: Word文档解析结果
            ocr_results: OCR结果列表
            
        Returns:
            结构比较结果
        """
        comparison = {
            "document": {
                "name": doc_result["file_name"],
                "paragraphs": len(doc_result["paragraphs"]),
                "tables": len(doc_result["tables"]),
                "images_referenced": len(doc_result["images"])
            },
            "attachments": {
                "total_files": len(ocr_results),
                "pdf_files": sum(1 for r in ocr_results if r["file_type"] == ".pdf"),
                "image_files": sum(1 for r in ocr_results if r["file_type"] != ".pdf")
            },
            "structure_match": True,
            "notes": []
        }
        
        # 检查附件数量是否匹配
        if len(doc_result["images"]) != len(ocr_results):
            comparison["structure_match"] = False
            comparison["notes"].append(
                f"文档中引用了 {len(doc_result['images'])} 个图片，"
                f"但提供了 {len(ocr_results)} 个附件"
            )
        
        return comparison
